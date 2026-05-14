import streamlit as st
import pandas as pd
from datetime import datetime
import io
from streamlit_gsheets import GSheetsConnection
from docx import Document  # 🔴 新增的 Word 處理套件

st.set_page_config(page_title="旅遊日誌與匯出", page_icon="📖", layout="wide")

# --- 0. 旅程動態對接邏輯 ---
target_sheet = st.session_state.get('active_trip_sheet', 'Exp_Yunnan2026')
target_name = st.session_state.get('active_trip_name', '2026 雲南探索 (預設)')

with st.sidebar:
    st.header("🎯 存檔目標")
    st.success(f"目前檢視旅程：\n\n**{target_name}**")
    st.caption(f"數據來源分頁：`{target_sheet}`")
    st.divider()

# --- 職人風格 CSS (保留全部原始設計) ---
st.markdown("""
    <style>
    h1 { border-bottom: 2px solid #b87333; padding-bottom: 10px; }
    .timeline-item { 
        border-left: 3px solid #b87333; 
        padding-left: 20px; 
        margin-bottom: 25px; 
        position: relative;
    }
    .timeline-item::before {
        content: '';
        position: absolute;
        left: -8px;
        top: 0;
        width: 13px;
        height: 13px;
        border-radius: 50%;
        background-color: #b87333;
    }
    .timeline-title { color: #e0e0e0; font-size: 1.2em; font-weight: bold; margin-bottom: 5px; }
    .timeline-details { color: #888; font-size: 0.95em; }
    .companion-tag { background-color: #333; padding: 6px 12px; border-radius: 15px; border-left: 3px solid #b87333; display: inline-block; font-size: 0.9em; margin: 2px;}
    </style>
""", unsafe_allow_html=True)

st.title(f"📖 旅程回顧與日誌匯出")
st.caption(f"當前結算專案：{target_name}")

# --- 1. 初始化與雲端資料撈取 ---
if 'companions_list' not in st.session_state:
    st.session_state.companions_list = []

footprints = st.session_state.get('footprint_data', [])
EXPECTED_COLUMNS = ["日期", "分類", "項目", "金額", "付款方式", "支付人", "來源"]

# 🔴 核心修正：從雲端精準分離「總預算」與「真實花費」
total_budget = 50000 # 預設值
total_spent = 0
expense_count = 0
has_expenses = False
df_expenses = pd.DataFrame(columns=EXPECTED_COLUMNS)

try:
    conn = st.connection("gsheets", type=GSheetsConnection)
    df_raw = conn.read(worksheet=target_sheet, ttl=0)
    
    if not df_raw.empty and "分類" in df_raw.columns:
        # 1. 獨立抓取雲端總預算 (分類 == '系統設定')
        budget_row = df_raw[df_raw['分類'] == '系統設定']
        if not budget_row.empty:
            total_budget = int(budget_row['金額'].values[0])
            
        # 2. 過濾出真實花費 (剔除系統設定)，解決金額暴增問題
        df_expenses = df_raw[df_raw['分類'] != '系統設定'].copy()
        
        if not df_expenses.empty and "金額" in df_expenses.columns:
            df_expenses["金額"] = pd.to_numeric(df_expenses["金額"], errors='coerce').fillna(0)
            total_spent = int(df_expenses["金額"].sum())
            expense_count = len(df_expenses)
            has_expenses = True

except Exception as e:
    st.error(f"❌ 雲端資料讀取失敗，請確認連線與分頁名稱 ({target_sheet}): {e}")

# --- 2. 頂部數據總覽 ---
st.header("📊 旅程結算總覽")
c1, c2, c3, c4 = st.columns(4)
c1.metric("總預算 (雲端同步)", f"NT$ {total_budget:,}")
c2.metric("總花費 (雲端同步)", f"NT$ {total_spent:,}")
c3.metric("剩餘預算", f"NT$ {total_budget - total_spent:,}")
c4.metric("足跡打卡 / 雲端記帳數", f"{len(footprints)} / {expense_count} 筆")

st.divider()

# --- 3. 版面分割：左側時間軸與表格，右側匯出功能 ---
col_left, col_right = st.columns([1.5, 1])

with col_left:
    st.subheader("🗺️ 足跡時間軸回顧")
    if footprints:
        sorted_footprints = sorted(footprints, key=lambda x: x.get("日期", "9999-99-99"))
        current_date = None
        for pt in sorted_footprints:
            pt_date = pt.get("日期", "未設定日期")
            if pt_date != current_date:
                st.markdown(f"### 📅 {pt_date}")
                current_date = pt_date
            
            st.markdown(f"""
            <div class="timeline-item">
                <div class="timeline-title">{pt.get('名稱')} <span style="font-size: 0.8em; color: gray;">({pt.get('類型')})</span></div>
                <div class="timeline-details">
                    ⏱️ 停留時間: {pt.get('停留時間', '未設定')} <br>
                    📝 筆記: {pt.get('描述', '無')}
                </div>
            </div>
            """, unsafe_allow_html=True)
    else:
        st.info("目前尚無足跡資料。")

    st.write("")
    st.subheader("💸 花費明細總覽 (含支付人資訊)")
    if has_expenses:
        display_cols = [c for c in EXPECTED_COLUMNS if c in df_expenses.columns]
        st.dataframe(df_expenses[display_cols].sort_values(by="日期", ascending=False), use_container_width=True, hide_index=True)
    else:
        st.info("目前雲端帳本尚無花費紀錄。")

with col_right:
    st.subheader("👥 旅程基本資訊")
    
    # 旅伴管理
    def add_companion():
        new_comp = st.session_state.new_comp_input.strip()
        if new_comp and new_comp not in st.session_state.companions_list:
            st.session_state.companions_list.append(new_comp)
        st.session_state.new_comp_input = ""

    st.text_input("➕ 新增同行旅伴", key="new_comp_input", on_change=add_companion, placeholder="輸入姓名後按 Enter...")
    
    if st.session_state.companions_list:
        for i, comp in enumerate(st.session_state.companions_list):
            c_tag, c_del = st.columns([4, 1])
            with c_tag:
                st.markdown(f"<div class='companion-tag'>👤 {comp}</div>", unsafe_allow_html=True)
            with c_del:
                if st.button("❌", key=f"del_comp_{i}"):
                    st.session_state.companions_list.pop(i)
                    st.rerun()
    
    st.divider()
    st.subheader("🖨️ 產生與匯出報告")
    
    # --- 4. 產生 Markdown 報告內容 ---
    report_md = f"# ✈️ 旅遊日誌：{target_name}\n\n"
    if st.session_state.companions_list:
        report_md += f"👥 **同行旅伴:** {', '.join(st.session_state.companions_list)}\n\n"
    
    report_md += f"## 📊 財務總結\n"
    report_md += f"- **設定總預算:** NT$ {total_budget:,}\n"
    report_md += f"- **實際總花費:** NT$ {total_spent:,}\n"
    report_md += f"- **結餘:** NT$ {total_budget - total_spent:,}\n\n"
    
    # --- 🔴 新增功能：產生 Word (Docx) 報告內容 ---
    def create_word_document():
        doc = Document()
        doc.add_heading(f"✈️ 旅遊日誌：{target_name}", 0)
        
        if st.session_state.companions_list:
            doc.add_heading("👥 同行旅伴", level=2)
            doc.add_paragraph(", ".join(st.session_state.companions_list))
            
        doc.add_heading("📊 財務總結", level=2)
        doc.add_paragraph(f"設定總預算: NT$ {total_budget:,}")
        doc.add_paragraph(f"實際總花費: NT$ {total_spent:,}")
        doc.add_paragraph(f"結餘: NT$ {total_budget - total_spent:,}")
        
        doc.add_heading("🗺️ 足跡時間軸回顧", level=2)
        if footprints:
            sorted_pt = sorted(footprints, key=lambda x: x.get("日期", "9999-99-99"))
            curr_d = None
            for pt in sorted_pt:
                if pt.get("日期") != curr_d:
                    doc.add_heading(f"📅 {pt.get('日期', '未設定日期')}", level=3)
                    curr_d = pt.get("日期")
                p = doc.add_paragraph()
                p.add_run(f"📍 {pt.get('名稱')} ({pt.get('類型')})\n").bold = True
                p.add_run(f"⏱️ 停留時間: {pt.get('停留時間', '未設定')}\n")
                p.add_run(f"📝 筆記: {pt.get('描述', '無')}")
        else:
            doc.add_paragraph("目前尚無足跡資料。")
            
        bio = io.BytesIO()
        doc.save(bio)
        return bio.getvalue()

    # --- 下載按鈕區 ---
    
    # 按鈕 1: Word 格式 (新功能)
    docx_data = create_word_document()
    st.download_button(
        label="📄 下載旅遊報告 (Word 格式)",
        data=docx_data,
        file_name=f"{target_sheet}_Log_{datetime.now().strftime('%Y%m%d')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        type="primary",
        use_container_width=True
    )
    
    # 按鈕 2: Markdown (回憶日誌)
    st.download_button(
        label="📥 下載回憶日誌 (Markdown)",
        data=report_md,
        file_name=f"{target_sheet}_Log_{datetime.now().strftime('%Y%m%d')}.md",
        mime="text/markdown",
        use_container_width=True
    )
    
    st.write("")
    
    # 按鈕 3: CSV (花費明細報表)
    if has_expenses:
        export_df = df_expenses.copy()
        if st.session_state.companions_list:
            export_df["旅伴紀錄"] = ", ".join(st.session_state.companions_list)
            
        csv_data = export_df.to_csv(index=False).encode('utf-8-sig')
        st.download_button(
            label="📊 匯出總花費明細 (CSV)",
            data=csv_data,
            file_name=f"{target_sheet}_Expenses_{datetime.now().strftime('%Y%m%d')}.csv",
            mime="text/csv",
            use_container_width=True
        )
    else:
        st.button("📊 匯出花費明細 (無資料)", disabled=True, use_container_width=True)

    st.caption("💡 提示：匯出的 CSV 包含所有的消費細節與**支付人**，直接用 Excel 開啟中文不會亂碼。")
